import docx
import json
import base64
import io
from PIL import Image
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
import numpy as np

def resize_and_autoorient(file, height,width):
    """Accepts a file bytes object and returns a file bytes object
    Resizes an image based on specifications in the config."""
    f = Image.open(file)
    #f = autoorient(f)
    f=f.resize((int(height),int(width)))
    # Create a bytes object to send in response
    img_file = io.BytesIO()
    f.save(img_file, format='JPEG')
    img_file.seek(0)
    return img_file

def create_word_table_from_json(doc, js,params={'image_properties':{'columns':['image'],'image_size':[125,80]}}):
    if not any(js): return doc
    
    ## Handling of input parameters
    header_row = True
    image_columns=[]
    col_width=[]
    image_size=[150,120]
    columns = list(js[0].keys())
    table_style,preset,paragraph_style, image_properties, paragraph = None,None,None,None,None
    if params:
        if "column_widths" in params.keys(): col_width = params['column_widths']
        if "preset" in params.keys(): preset = params['preset']
        if "paragraph" in params.keys():paragraph = params['paragraph']
        if "table_style" in params.keys(): table_style = params['table_style']
        if "image_properties" in params.keys(): image_properties = params['image_properties']
        if "columns" in params.keys(): columns = params['columns']
        if "header_row" in params.keys(): header_row = params['header_row']
        
    if image_properties:
        if "columns" in image_properties.keys():image_columns = image_properties['columns']
        if "image_size" in image_properties.keys(): image_size = image_properties['image_size']

    if paragraph:
        if "paragraph_style" in paragraph.keys(): paragraph_style = paragraph['paragraph_style']
        else: 
            styles = doc.styles
            randstyle = 'Customtablestyle'+str(np.random.randint(1, 1000001))
            if randstyle in styles: randstyle = 'Customtablestyle'+str(np.random.randint(1000001, 2000001))
            style = styles.add_style(randstyle, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = styles['Normal']
            if "font" in paragraph_style.keys():
                fontjs = paragraph_style['font']
                font=style.font
                if "size" in fontjs.keys():
                    font.size = Pt(int(fontjs['size']))
                if "name" in fontjs.keys():
                    font.name = fontjs['name']
                if "bold" in fontjs.keys():
                    font.bold = fontjs['bold']
                if "italic" in fontjs.keys():
                    font.italic = fontjs['italic']
                    
                    
                    
        
    table = doc.add_table(rows=1, cols = len(columns))
    
    if table_style:
        table.style=table_style
    if True:
        for item in js:
            ## Sätt rubriker till tabellen
            if header_row==True:
                header_row = table.rows[0].cells
                for i in range(len(columns)):
                    p = header_row[i].paragraphs[0]
                    p.text = columns[i]
                    p.runs[0].bold=True
        
        
            row = table.add_row().cells
            for i, column in enumerate(columns):
                if column in image_columns:
                    file = io.BytesIO(base64.b64decode(item[column]))
                    file.seek(0)
                    file = resize_and_autoorient(file,*image_size)
                    p=row[i].paragraphs[0]
                    p.style.paragraph_format.keep_with_next=True
                    run = p.add_run()
                    picture = run.add_picture(file)
                    continue
                    
                row[i].text = str(item[column])
                p = row[i].paragraphs[0]
                if paragraph_style: p.style = paragraph_style
                p.paragraph_format.keep_with_next=True
                
        if col_width:
            for i, column in enumerate(table.columns):
                width = col_width[i]
                for cell in column.cells:
                    cell.width = Inches(width)
                    
    return doc


if __name__ == '__main__':
    doc = docx.Document()
    
    
    with open('image.jpg', 'rb') as f:
        img = base64.b64encode(f.read())
        
    params={
            'image_properties':
                {
                    'columns':['image'],
                    'image_size':[80,80]
                    },
            'columns': ['name','age','email','image'],
            'column_widths':[0.5,0.8,1,6],
            'table_style': 'Table Grid'
                }
    js = [
    {
        "id": 1,
        "name": "John Doe",
        "age": 30,
        "email": "john.doe@example.com",
        "image": img
    },
    {
        "id": 2,
        "name": "Jane Smith",
        "age": 25,
        "email": "jane.smith@example.com",
        "image": img
    },
    {
        "id": 3,
        "name": "Alex Johnson",
        "age": 35,
        "email": "alex.johnson@example.com",
        "image": img
    }
]
    create_word_table_from_json(doc, js, params).save('001.docx')