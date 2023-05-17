import json
import os
import requests
from functions.downloader.sharepoint import get_sharepoint_access_headers_through_client_id,get_item_based_on_id, get_all_items, get_by_url, Download_icon
import datetime
import io
import mailmerge
from PIL import Image
import docx
from docx.shared import Pt, Inches, RGBColor
from Office_helper_functions.Word.form_field import set_checkbox_value, compress_word_file
from Office_helper_functions.Image.Image_operations import resize_and_autoorient
from functions.Word.add_image_to_zipfile import add_icon_to_word_file
import base64
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT, WD_UNDERLINE
from flask import abort

icons = {
         "Green Landscaping AB":"https://greenlandscapingmalmo.sharepoint.com/sites/Funktionskontrolllekplatsdemo/_api/web/GetFileByServerRelativeUrl('/sites/Funktionskontrolllekplatsdemo/Delade Dokument/Protokoll lekplats besiktning demo/Template/Loggor/GLAB_png.png')/$value",
         "Green Landscaping Skåne AB":"https://greenlandscapingmalmo.sharepoint.com/sites/Funktionskontrolllekplatsdemo/_api/web/GetFileByServerRelativeUrl('/sites/Funktionskontrolllekplatsdemo/Delade Dokument/Protokoll lekplats besiktning demo/Template/Loggor/GLAB_png.png')/$value",
         "Svensk Markservice":"https://greenlandscapingmalmo.sharepoint.com/sites/Funktionskontrolllekplatsdemo/_api/web/GetFileByServerRelativeUrl('/sites/Funktionskontrolllekplatsdemo/Delade Dokument/Protokoll lekplats besiktning demo/Template/Loggor/SvMS_png.png')/$value",
         "Green Landscaping Malmö AB":"https://greenlandscapingmalmo.sharepoint.com/sites/Funktionskontrolllekplatsdemo/_api/web/GetFileByServerRelativeUrl('/sites/Funktionskontrolllekplatsdemo/Delade Dokument/Protokoll lekplats besiktning demo/Template/Loggor/GL_Malmo_png.png')/$value",
         "Markservice Stockholm AB":"https://greenlandscapingmalmo.sharepoint.com/sites/Funktionskontrolllekplatsdemo/_api/web/GetFileByServerRelativeUrl('/sites/Funktionskontrolllekplatsdemo/Delade Dokument/Protokoll lekplats besiktning demo/Template/Loggor/markservice_sthlm_png.png')/$value",
         "PARK i Syd AB":"https://greenlandscapingmalmo.sharepoint.com/sites/Funktionskontrolllekplatsdemo/_api/web/GetFileByServerRelativeUrl('/sites/Funktionskontrolllekplatsdemo/Delade Dokument/Protokoll lekplats besiktning demo/Template/Loggor/PARK green 4F.png')/$value",
         "Tranemo Trädgårdstjänst AB":"https://greenlandscapingmalmo.sharepoint.com/sites/Funktionskontrolllekplatsdemo/_api/web/GetFileByServerRelativeUrl('/sites/Funktionskontrolllekplatsdemo/Delade Dokument/Protokoll lekplats besiktning demo/Template/Loggor/TTJ_logo_ute_webb.png')/$value",
         "Grön Stad":"https://greenlandscapingmalmo.sharepoint.com/sites/Funktionskontrolllekplatsdemo/_api/web/GetFileByServerRelativeUrl('/sites/Funktionskontrolllekplatsdemo/Delade Dokument/Protokoll lekplats besiktning demo/Template/Loggor/Gron_Stad_logo_tag_RGB_270px.png')/$value"
         }




def get_cert_no(site,certlist, fitness):
    items = get_all_items(site=site, list_=certlist)
    itemsurl = items["d"]["Items"]["__deferred"]["uri"]
    items = requests.get(itemsurl)


def create_protocol(site, lista, js):

    
    if 'Certifikatinfo' not in js.keys():js = js['body']
    certifikatjs = js['Certifikatinfo']['value'][0]
    
    for item in js['Items']['value']:
        if 'Adress' not in item.keys():
            item['Adress'] = ' '
    trigger = js['Trigger']
    
    js1 = js["Items"]['value'][0]
    if not any(certifikatjs): js1['Certnr'] = 'saknas'
    js1['Informationsskylt'] = ['Finns' if js1['Informationsskylt'] else 'Saknas på ett eller flera redskap'][0]
    js1['Anv_x00e4_ndarinformation'] = ['Finns' if js1['Anv_x00e4_ndarinformation'] else 'Saknas på ett eller flera redskap'][0]
    js1['M_x00e4_rkningavredskap_x002f_ty'] = ['Finns' if js1['M_x00e4_rkningavredskap_x002f_ty'] else 'Saknas på ett eller flera redskap'][0]
    js1['I_bed'] = ['-' if js1['Informationsskylt']=='Finns' else 'C'][0]
    js1['A_bed'] = ['-' if js1['Anv_x00e4_ndarinformation']=='Finns' else 'C'][0]
    js1['M_bed'] = ['-' if js1['M_x00e4_rkningavredskap_x002f_ty']=='Finns' else 'C'][0]
    if "Telefonnummer" not in js1.keys(): js1["Telefonnummer"] = ''

    js1['Hemsida'] = certifikatjs['Hemsida']
    js1['Email'] = js1['Author']['Email']
    js1['Dagensdatum'] = datetime.datetime.strftime(datetime.datetime.now(), "%Y-%m-%d")
    js1['Bolag'] = certifikatjs['Bolag']['Value']
    js1["Certifieringstext2"] = """Klagomål

I det fall Ni har synpunkter på utförd säkerhetsbesiktning kan Ni Kontakta: 

SERENO Certifiering AB
Box 5604

114 86  STOCKHOLM 
Telefon: 08-556 953 30 
"""
    js1["Certifieringstext"] = """Innehar Certifikat nr Click or tap here to enter text. utfärdat av
SERENO Certifiering AB
Platser för motion eller annan utevistelse
Samt av fitnessutrustning
"""
    
    doc = populate_template(js1,certifikatjs,js,trigger)
    return doc
    
    
def populate_template(js1, certifikatjs, js, trigger):
    
    if trigger['DigitalsignaturLekplats'] or trigger['DigitalsignaturUtegym']:
        js1['Digital signatur'] = "Härmed intygas att besiktningen utförts enligt gällande regler."
        js1['Digital signatur 2'] = "Digitalt signerad av "+trigger['Author']['DisplayName']+', '+ trigger['Created'].split('T')[0]
    if 'Digital signatur' in js1.keys() and 'Digital signatur 2' in js1.keys():
        js1['digsign'] = js1['Digital signatur']
        js1['digsign2'] = js1['Digital signatur 2']
    if "Telefonnummer" not in certifikatjs.keys(): abort(400, message="Inget telefonnummer för besiktningsman.")
    js1["Besmantelefonnummer"] = certifikatjs["Telefonnummer"]
    js1["Adresstillprotokoll"] = certifikatjs['Adresstillprotokoll']
    js1['Created'] = js1['Created'].split('T')[0]
    if not 'Certnr' in certifikatjs.keys(): certifikatjs['Certnr'] = 'saknas'
    js1['Certnr'] = certifikatjs['Certnr']
    
    if js1["Certnr"].lower() == 'saknas' and js1['Fitnessbesiktning']:
        doc = mailmerge.MailMerge(os.path.join(os.path.dirname(__file__), 'Fitness mall ej cert.docx'))
    elif js1["Certnr"].lower() != 'saknas' and js1['Fitnessbesiktning']:
        doc = mailmerge.MailMerge(os.path.join(os.path.dirname(__file__), 'Fitness mall cert.docx'))
    elif js1["Certnr"].lower() != 'saknas' and not js1['Fitnessbesiktning']:
        doc = mailmerge.MailMerge(os.path.join(os.path.dirname(__file__), 'Lekplatsbesiktning mall cert.docx'))
    elif js1["Certnr"].lower() == 'saknas' and not js1['Fitnessbesiktning']:
        doc = mailmerge.MailMerge(os.path.join(os.path.dirname(__file__), 'Lekplatsbesiktning mall ej cert.docx'))
 

    js1["H_x00e4_nvisningsskylt"] = "Ja" if js1["H_x00e4_nvisningsskylt"]==True else "Nej"
    js1["Rutinerolyckor_x002f_incidenter"] = "Ja" if js1["Rutinerolyckor_x002f_incidenter"]==True else "Nej"
    js1["Rutinerf_x00f6_runderh_x00e5_ll"] = "Ja" if js1["Rutinerf_x00f6_runderh_x00e5_ll"]==True else "Nej"
    js1["Leverant_x00f6_rensanvisningar"] = "Ja" if js1["Leverant_x00f6_rensanvisningar"]==True else "Nej"
    js1["Sparandeavdokument"] = "Ja" if js1["Sparandeavdokument"]==True else "Nej"
    
    mergefields = {key:value for key,value in js1.items() if key in [x for x in doc.get_merge_fields()]}


    doc.merge(**mergefields)
    
    if js1['Bolag'] not in icons.keys(): bolag = "Green Landscaping AB"
    else: bolag = js1['Bolag']
    icon_file = requests.get(icons[bolag], headers=get_sharepoint_access_headers_through_client_id()).content
    ### TYPE CHANGE: type(doc) = docx.Document object from here.
    doc = change_icon_in_header(doc, js1, icon_file, imagepath="word/media/image1.png")
    if js1['Typavbesiktning']['Value'] == "Installationsbesiktning":
        file = io.BytesIO()
        doc.save(file)
        file.seek(0)
        doc = docx.Document(set_checkbox_value(file.read(), 1,1))
    else: 
        file = io.BytesIO()
        doc.save(file)
        file.seek(0)
        doc = docx.Document(set_checkbox_value(file.read(), 1,0)) #Årlig besiktning
    
    
    if not js1["Fitnessbesiktning"]:

        match js1["Fysiskomfattning"]['Value']:
            case "Endast lekredskap":
                file = io.BytesIO()
                doc.save(file)
                file.seek(0)
                doc = docx.Document(set_checkbox_value(file.read(), 1,2)) # Årlig besiktning
            case "Lekplats inklusive":
                file = io.BytesIO()
                doc.save(file)
                file.seek(0)
                doc = docx.Document(set_checkbox_value(file.read(), 1,3)) # Lekplats inklusive
            case "Lekplats inkl staket mm":
                file = io.BytesIO()
                doc.save(file)
                file.seek(0)
                doc = docx.Document(set_checkbox_value(file.read(), 1,3)) # Lekplats inklusive
            case "Område på karta":
                file = io.BytesIO()
                doc.save(file)
                file.seek(0)
                doc = docx.Document(set_checkbox_value(file.read(), 1,4)) # Område på karta

        match js1["Drift_x002d_ochunderh_x00e5_llsp"]:
            case True:
                file = io.BytesIO()
                doc.save(file)
                file.seek(0)
                doc = docx.Document(set_checkbox_value(file.read(), 1,5)) # Underhållsplan True
            case False:
                file = io.BytesIO()
                doc.save(file)
                file.seek(0)
                doc = docx.Document(set_checkbox_value(file.read(), 1,6)) # Underhållsplan False
    st = [style.type for style in doc.styles if style.name == 'Heading 1'][0]
    bighead = doc.styles.add_style('Big heading', st)
    bighead.font.size = Pt(22)
    bighead.font.color.rgb = RGBColor(96,167,48)
    smallhead = doc.styles.add_style('Small heading', st)
    smallhead.font.size = Pt(16)
    smallhead.font.color.rgb = RGBColor(96,167,48)
    imgp = doc.styles.add_style('imgp', st)
    imgp.font.size = Pt(11)
    imgp.font.color.rgb = RGBColor(96,167,48)

    doc.add_page_break()

    add_översiktsbild(doc,js)
    add_utrustning(doc,js)

    if len(js['Utrustning']) >7 and len(js['Utrustning']) <= 12: doc.add_page_break()
    # add_page_break(doc)
    add_anmärkningar(doc,js)
    # add_page_break(doc)
    add_underlag(doc,js)
    if any(js['Staket']):
        add_grindar(doc,js)
    if any(js['Brunnar']):
        add_brunnar(doc,js)
    else:
        js['Brunnar'] = [{'Items':{
            '{HasAttachments}':False,
            'Anm_x00e4_rkning':'-',
            'Status':{'Value':'Ej kontrollerade.'}
            }, 'Images':''}]
        add_brunnar(doc,js)
    #doc.save(os.path.join(os.path.dirname(__file__), 'steg_1.docx'), )
    return doc



def add_översiktsbild(doc,js):
    if js['Trigger']['Sida2']:
        image = io.BytesIO(base64.b64decode(js['Översiktsbild']))
        image.seek(0)
        h=doc.add_heading('Översiktsbild av lekplatsen', 0)
        h.style = 'Big heading'
        img = resize_and_autoorient(image,300,320)
        table = doc.add_table(rows=1,cols=1)
        row = table.add_row().cells
        row[0].paragraphs[0].add_run().add_picture(img)
    return None


def add_utrustning(doc,js):
    # Add the table for the Utrustnings items:
    h = doc.add_heading('Produktbeskrivning')
    h.style = 'Big heading'
    subheading = doc.styles.add_style('subheading', h.style.type)
    subheading.font.size = Pt(16)
    subheading.font.color.rgb = RGBColor(96,167,48)
    subheading2 = doc.styles.add_style('subheading2', h.style.type)
    subheading2.font.size = Pt(14)
    subheading2.font.color.rgb = RGBColor(96,167,48)
    small = doc.styles.add_style('small', h.style.type)
    small.font.size = Pt(10)
    small.font.italic = True
    vsmall= doc.styles.add_style('vsmall', h.style.type)
    vsmall.font.size = Pt(8)
    bold = doc.styles.add_style('bold', h.style.type)
    bold.font.size = Pt(10)
    bold.font.bold = True
    # h.style.font.color.rgb = RGBColor(100,200,100)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Grid Table 1 Light'
    row = table.rows[0].cells
    row[0].text = 'Nr'
    row[0].width = Inches(0.2)
    row[0].paragraphs[0].paragraph_format.keep_with_next=True
    row[1].text = "Produkt"
    row[1].paragraphs[0].paragraph_format.keep_with_next=True
    row[2].text = "Tillverkare/artnr"
    row[2].paragraphs[0].paragraph_format.keep_with_next=True
    row[3].text = "Årtal"
    row[3].paragraphs[0].paragraph_format.keep_with_next=True
    row[4].text = "Bild nr"
    row[4].paragraphs[0].paragraph_format.keep_with_next=True
    
    for i, item in enumerate(js["Utrustning"]):
        row = table.add_row()
        row=row.cells
        row[0].text = str(i+1)
        row[0].paragraphs[0].paragraph_format.keep_with_next=True
        row[0].width = Inches(0.2)
        
        if "Utegymredskap" in item['Items'].keys() and "Utrustning" not in item['Items'].keys(): row[1].text = item['Items']['Utegymredskap']['Value']
        else : row[1].text = item['Items']['Utrustning']['Value']
        row[1].paragraphs[0].paragraph_format.keep_with_next=True
        if "Tillverkare_x002f_artnr" in item['Items'].keys(): 
            row[2].text = item['Items']['Tillverkare_x002f_artnr']
        else: row[2].text = '-'
        row[2].paragraphs[0].paragraph_format.keep_with_next=True

        if "OData__x00c5_rtal" not in item["Items"].keys(): item['Items']['OData__x00c5_rtal'] = 'Saknas'
        row[3].text = item['Items']['OData__x00c5_rtal']
        row[3].paragraphs[0].paragraph_format.keep_with_next=True
        row[4].text = "Bild: "+str(i+1)
        if i != len(js['Utrustning'])-1:row[4].paragraphs[0].paragraph_format.keep_with_next=True
    
    for cell in table.columns[0].cells:
        cell.width = Inches(0.5)
    for cell in table.columns[1].cells:
        cell.width = Inches(2.5)
    for cell in table.columns[2].cells:
        cell.width = Inches(2.5)
    for cell in table.columns[3].cells:
        cell.width = Inches(0.7)
    for cell in table.columns[4].cells:
        cell.width = Inches(0.8)
    hh = doc.add_heading('Besiktningsresultat', 0)
    hh.style = 'Big heading'
    hh.style.paragraph_format.keep_with_next=True
    
    
    
    
    

    table = doc.add_table(rows=0,cols=3)
    index1 = 1
    imagedict = {}
    for item in js['Utrustning']:
        for item2 in item['Image']:
            imagedict[item2['content']] = index1
            
        index1+=1
    index0 = 0
    row = table.add_row().cells
    for img, index in imagedict.items():
        pt = row[index0].add_paragraph()

        pt.style = 'Small heading'

        pt.style.paragraph_format.keep_with_next=True
        pt.text = 'Bild: ' + str(index)
        file = io.BytesIO(base64.b64decode(img))
        file.seek(0)
        file = resize_and_autoorient(file,120,120)
        p = row[index0].paragraphs[0]
        p.add_run().add_picture(file)
        index0 +=1
        if index0 ==3:
            row = table.add_row().cells
            index0 = 0
    if js['Items']['value'][0]['Typavbesiktning']['Value'] == 'Installationsbesiktning':
        
        for i, item in enumerate(js['Utrustning']):
            if "Utegymredskap" in item['Items'].keys() and "Utrustning" not in item['Items'].keys(): 
                print("Hello")
                Produkt = item['Items']['Utegymredskap']['Value']
                item['Items']['Utrustning'] = {'Value':item['Items']['Utegymredskap']}
            elif 'Utrustning' in item['Items'].keys(): 
                Produkt = item['Items']['Utrustning']['Value']
            else: 
                Produkt = 'saknas'
                item['Items']['Utrustning'] = {'Value':'saknas'}
                
            

            # table = doc.add_table(rows=1, cols=5)
            # table.style = 'Grid Table 1 Light'
            # row = table.rows[0].cells
            # row[0].text = 'Nr'
            # row[0].width = Inches(0.2)
            # row[1].text = "Produkt"
            # row[2].text = "Tillverkare/artnr"
            # row[3].text = "Årtal"
            # row[4].text = "Bild nr"
    # while index1 < len(js['Utrustning']):
    #     index = 0
    #     while index < len(js['Utrustning'][index1]['Image']):
    #         row = table.add_row().cells
    #         for i in range(0,3):
    #             if index1 == len(js['Utrustning']):break
    #             if index == len(js['Utrustning'][index1]['Image']): continue
    #             p= row[i].add_paragraph()
    #             p.text = 'Bild: '+str(index1+1)
    #             file = io.BytesIO(base64.b64decode(item['Image'][index]['content']))
    #             file.seek(0)
    #             file = resize_and_autoorient(file,120,120)
    #             row[i].add_paragraph().add_run().add_picture(file)
    #             index+=1
    #             index1+=1
            
    # index = 0
    # table = doc.add_table(rows=1, cols=3)
    # while index < len(images):
    #     row = table.add_row().cells
    #     for i in range(0,3):
            
    #         file = io.BytesIO(base64.b64decode(js['Utrustning'][index]['Image'][0]['content']))
    #         file.seek(0)
    #         file = resize_and_autoorient(file,120,120)
    #         row[i].add_paragraph().text = "Bild: "+str(index+1)
    #         row[i].add_paragraph().add_run().add_picture(file)
    #         index +=1
    #         if index == len(images):break
    return None

def add_underlag(doc,js):
    doc.add_page_break()
    doc.add_paragraph()
    hh = doc.add_heading('Stötdämpande underlag:', 0)
    hh.style = 'Big heading'
    #images = [img['Image'][0] for img in js['Anmärkningar']]
    hh.paragraph_format.keep_with_next=True
    if not any(js['Underlag']):
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Grid Table Light'
        table.style.paragraph_format.keep_with_next = True
        row = table.rows[0].cells
        row[0].text = 'Inga kommentarer gällande underlag'
        row[0].paragraphs[0].paragraph_format.keep_with_next=True
        row[0].style = 'vsmall'
        row[1].text = '-'
        row[1].paragraphs[0].paragraph_format.keep_with_next=True
        p = doc.add_paragraph()
        p.text = "Enligt SS-EN 1176-1:4.2.8.5"
        p.style = 'small'
        for cell in table.columns[0].cells:
            cell.width = Inches(6)
        for cell in table.columns[1].cells:
            cell.width = Inches(0.4)
        return None
    count=0
    for i, item in enumerate(js['Underlag']):
        p = doc.add_paragraph()
        if 'Utrustning' not in item.keys():
            if 'Kommentar' in item.keys():
                item['Utrustning'] = item['Kommentar']
            else:
                item['Utrustning'] = 'Utrustning'
                item['Kommentar'] = ['-']
        if not any([i+1 for i, utr in enumerate(js['Utrustning']) if utr['Items']['ID'] == item['UtrustningsID']]):
            continue
        count+=1
        p.text = 'Produkt '+str([i+1 for i, utr in enumerate(js['Utrustning']) if utr['Items']['ID'] == item['UtrustningsID']][0])+':' + item['Utrustning']
        p.style = 'bold'
        p.paragraph_format.keep_with_next = True
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Grid Table Light'
        table.style.paragraph_format.keep_with_next = True
        row = table.rows[0].cells
        row[0].text = item['Kommentar']
        row[0].paragraphs[0].paragraph_format.keep_with_next=True
        row[0].style = 'vsmall'
        row[1].text = item['Bed_x00f6_mning']['Value']
        row[1].paragraphs[0].paragraph_format.keep_with_next=True
        p = doc.add_paragraph()
        p.text = "Enligt SS-EN 1176-1:4.2.8.5"
        p.style = 'small'
        for cell in table.columns[0].cells:
            cell.width = Inches(6)
        for cell in table.columns[1].cells:
            cell.width = Inches(0.4)
    return None
    
def add_anmärkningar(doc, js):
    doc.add_page_break()
    hh = doc.add_heading('Anmärkningar:', 0)
    hh.style = 'Big heading'
    hh.paragraph_format.keep_with_next = True
    for i,utrustning in enumerate(js['Utrustning']):
        utrustning=utrustning['Items']
        anmärkningar = [anmärkning for anmärkning in js['Anmärkningar'] if anmärkning['Items']['UtrustningsID'] == utrustning['ID']]
        #print(utrustning.keys())
        if 'Utrustning' not in utrustning.keys():
            utrustning['Utrustning'] =  {'Value':utrustning['Utegymredskap']['Value']}
        h = doc.add_heading('Produkt '+str(i+1)+', '+ utrustning['Utrustning']['Value'], 0)
        h.style= 'subheading'
        h.paragraph_format.keep_with_next = True
        h.runs[0].bold=True
        
                    # LOOP FÖR ATT LÄGGA TILL BILDER
        # print(any(anmärkningar))
        # if any(anmärkningar):
        #     ph = doc.add_paragraph()
        #     ph.text = "Anmärkningar"
        #     ph.style = 'subheading2'
        #     ph.style.paragraph_format.keep_with_next=True
        for anmärkning in anmärkningar:
            print("Hello")
            if anmärkning['Items']['{HasAttachments}']:
                table=doc.add_table(rows=0, cols=4)
                index = 0
                table.style.paragraph_format.keep_together = True
                while index < len(anmärkning['Image']):
                    row0 = table.add_row()
                    row =row0.cells
                    for i in range(0,4):
                        if index == len(anmärkning['Image']): continue
                        row[i]
                        file = io.BytesIO(base64.b64decode(anmärkning['Image'][index]['content']))
                        file.seek(0)
                        file = resize_and_autoorient(file,100,100)
                        p=row[i].paragraphs[0]
                        p.style= 'imgp'
                        p.paragraph_format.keep_with_next=True
                        run = p.add_run()
                        picture = run.add_picture(file)
                        index +=1
                        
                        #  LÄGG TILL ANMÄRKNINGAR
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Grid Table Light'
            table.style.paragraph_format.keep_with_next = True
            row = table.rows[0].cells
            row[0].text = anmärkning['Items']['Kommentar']
            row[0].paragraphs[0].paragraph_format.keep_with_next=True
            row[1].text = anmärkning['Items']['Bed_x00f6_mning']['Value']
            row[1].paragraphs[0].paragraph_format.keep_with_next=True
            for cell in table.columns[0].cells:
                cell.width = Inches(6)
            for cell in table.columns[1].cells:
                cell.width = Inches(0.4)
                # Standard under tabell
            p1 = doc.add_paragraph()
            p1.text = anmärkning['Items']['Utrustningstyp']['Value']
            p1.style = 'small'
        if not any(anmärkningar):
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Grid Table Light'
            table.style.paragraph_format.keep_with_next = True
            row = table.rows[0].cells
            row[0].text = "Inga anmärkningar funna vid besiktningstillfället"
            row[0].paragraphs[0].paragraph_format.keep_with_next=True
            row[1].text = "-"
            row[1].paragraphs[0].paragraph_format.keep_with_next=True
            for cell in table.columns[0].cells:
                cell.width = Inches(6)
            for cell in table.columns[1].cells:
                cell.width = Inches(0.4)
                # Standard under tabell
            p1 = doc.add_paragraph()
            p1.text = "SS-EN 1176-1177 alt 16630"
            p1.style = 'small'
            
            # LÄGG TILL MONTERING OVAN OCH UNDER MARK
        if False:
            if 'Montering_ovan_mark' not in utrustning.keys(): utrustning['Montering_ovan_mark'] = '-'
            if 'Montering_under_mark' not in utrustning.keys(): utrustning['Montering_under_mark'] = '-'
            if 'Montering_ovan_bed' not in utrustning.keys():
                utrustning['Montering_ovan_bed'] = {'Value':'-'}
            if 'Montering_under_bed' not in utrustning.keys():
                utrustning['Montering_under_bed']={'Value':'-'}
            if js['Items']['value'][0]['Typavbesiktning']['Value']=='Installationsbesiktning':
                ph = doc.add_paragraph()
                ph.text = "Montering ovan mark"
                ph.style = 'subheading2'
                ph.style.paragraph_format.keep_with_next=True
                
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Grid Table Light'
                table.style.paragraph_format.keep_with_next = True
                row = table.rows[0].cells
                row[0].text = utrustning['Montering_ovan_mark']
                row[0].paragraphs[0].paragraph_format.keep_with_next=True
                row[1].text = utrustning['Montering_ovan_bed']['Value']
                row[1].paragraphs[0].paragraph_format.keep_with_next=True
                for cell in table.columns[0].cells:
                    cell.width = Inches(6)
                for cell in table.columns[1].cells:
                    cell.width = Inches(0.4)
                p = doc.add_paragraph()
                p.text = "Enligt SS EN 1176-1:6.2.2"
                p.style = 'small'
                p.paragraph_format.keep_with_next=True
                    
                    
                ph = doc.add_paragraph()
                ph.text = "Montering under mark"
                ph.style = 'subheading2'
                ph.style.paragraph_format.keep_with_next=True
                
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Grid Table Light'
                table.style.paragraph_format.keep_with_next = True
                row = table.rows[0].cells
                row[0].text = utrustning['Montering_under_mark']
                row[0].paragraphs[0].paragraph_format.keep_with_next=True
                row[1].text = utrustning['Montering_under_bed']['Value']
                row[1].paragraphs[0].paragraph_format.keep_with_next=True
                for cell in table.columns[0].cells:
                    cell.width = Inches(6)
                for cell in table.columns[1].cells:
                    cell.width = Inches(0.4)
                p = doc.add_paragraph()
                p.text = "Enligt SS EN 1176-1:6.2.2"
                p.style = 'small'
            
    
def __add_anmärkningar_deprecated(doc, js):
    doc.add_paragraph()
    hh = doc.add_heading('Anmärkningar:', 0)
    hh.style = 'Big heading'
    hh.style.paragraph_format.keep_with_next = True
    #images = [img['Image'][0] for img in js['Anmärkningar']]

    for i, item in enumerate(js['Anmärkningar']):

        utrustningslista = [(utrustning['Items']['Montering_under_mark'], utrustning['Items']['Montering_ovan_mark']) for utrustning in js['Utrustning'] if utrustning['Items']['ID'] == item['Items']['UtrustningsID'] and 'Montering_ovan_mark' in utrustning['Items'].keys()]
        if False: #(utrustningslista):
            montering_under_mark, montering_ovan_mark = utrustningslista[0]
            item['montering_under_mark'] = montering_under_mark
            item['montering_ovan_mark'] = montering_ovan_mark
            
        if 'Utrustningstyp0' or 'Utrustningstyp' in item['Items'].keys(): 
            utrustning = [items['Items']['Utrustning']['Value'] for items in js['Utrustning'] if items['Items']['ID'] == item['Items']['UtrustningsID']][0]
            h = doc.add_heading('Produkt '+str(i+1)+', '+ utrustning, 0)
            h.style= 'subheading'
            h.paragraph_format.keep_with_next = True
        else: 
            h = doc.add_heading('Produkt '+str(i+1)+', '+'Gymutrustning')
            h.style= 'subheading'
            h.paragraph_format.keep_with_next = True
        
        if item['Items']['{HasAttachments}']:
            table=doc.add_table(rows=0, cols=4)
            index = 0
            table.style.paragraph_format.keep_together = True
            while index < len(item['Image']):
                row0 = table.add_row()
                row =row0.cells
                for i in range(0,4):
                    if index == len(item['Image']): continue
                    row[i]
                    file = io.BytesIO(base64.b64decode(item['Image'][index]['content']))
                    file.seek(0)
                    file = resize_and_autoorient(file,100,100)
                    p=row[i].add_paragraph()
                    p.style= 'imgp'
                    # p.style.paragraph_format.keep_with_next=True
                    run = p.add_run()
                    picture = run.add_picture(file)
                    index +=1
                    
                    
        
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Grid Table Light'
        table.style.paragraph_format.keep_with_next = True
        row = table.rows[0].cells
        if "Kommentar" not in item['Items'].keys():item['Items']['Kommentar'] = '-'
        row[0].text = item['Items']['Kommentar']
        row[1].text = item['Items']['Bed_x00f6_mning']['Value']
        for cell in table.columns[0].cells:
            cell.width = Inches(6)
        for cell in table.columns[1].cells:
            cell.width = Inches(0.4)
        
        p1 = doc.add_paragraph()
        p1.text = item['Items']['Utrustningstyp']['Value']
        p1.style = 'small'
        
        # smallhead = doc.styles.add_style('Underovanmark', doc.styles)
        # smallhead.font.size = Pt(14)
        # smallhead.font.color.rgb = RGBColor(100,200,100)
        if False:#any(utrustningslista):
            ph = doc.add_paragraph()
            ph.text = "Montering ovan mark"
            ph.style = 'subheading2'
            ph.style.paragraph_format.keep_with_next=True
            
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Grid Table Light'
            table.style.paragraph_format.keep_with_next = True
            row = table.rows[0].cells
            row[0].text = montering_ovan_mark
            row[1].text = '-'
            for cell in table.columns[0].cells:
                cell.width = Inches(6)
            for cell in table.columns[1].cells:
                cell.width = Inches(0.4)
            p = doc.add_paragraph()
            p.text = "Enligt SS EN 1176-1:6.2.2"
            p.style = 'small'
                
                
            ph = doc.add_paragraph()
            ph.text = "Montering under mark"
            ph.style = 'subheading2'
            ph.style.paragraph_format.keep_with_next=True
            
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Grid Table Light'
            table.style.paragraph_format.keep_with_next = True
            row = table.rows[0].cells
            row[0].text = montering_under_mark
            row[1].text = '-'
            for cell in table.columns[0].cells:
                cell.width = Inches(6)
            for cell in table.columns[1].cells:
                cell.width = Inches(0.4)
            p = doc.add_paragraph()
            p.text = "Enligt SS EN 1176-1:6.2.2"
            p.style = 'small'

            
            
        # row0 = table.rows[0].cells
        # h = doc.add_heading('Produkt '+str(i+1)+', '+ [item['Items']['Utrustningstyp0'] if 'Utrustningstyp0' in item['Items'].keys() else 'Gymredskap'][0], 0)
        # h.style= 'subheading'
        # if item['Items']['{HasAttachments}']:
        #     if item['Image'][0]['content']:
        #         img = item['Image'][0]['content']
        #         file = io.BytesIO(base64.b64decode(img))
        #         file.seek(0)
        #         img = resize_and_autoorient(file, 80,80)
        #         p = doc.add_paragraph()
        #         p.add_run().add_picture(img)
        # table = doc.add_table(rows=1, cols=2)
        # table.style = 'Grid Table Light'
        # row = table.rows[0].cells
        # row[0].text = item['Items']['Kommentar']
        # row[1].text = item['Items']['Bed_x00f6_mning']['Value']
        # for cell in table.columns[0].cells:
        #     cell.width = Inches(6)
        # for cell in table.columns[1].cells:
        #     cell.width = Inches(0.4)
        # p1 = doc.add_paragraph()
        # p1.text = item['Items']['Utrustningstyp']['Value']
        # p1.style = 'small'

    return None

def add_grindar(doc, js):
    doc.add_paragraph()
    hh = doc.add_heading('Grindar och staket', 0)
    hh.style = 'Big heading'
    hh.paragraph_format.keep_with_next=True
    for i, item in enumerate(js['Staket']):
        
        if item['Items']['{HasAttachments}']:
            table = doc.add_table(rows=0, cols=4)
            index = 0
            table.style.paragraph_format.keep_with_next = True
            while index < len(item['Images']):
                row = table.add_row().cells
                for i in range(0,4):
                    if index == len(item['Images']): continue
                    row[i]
                    file = io.BytesIO(base64.b64decode(item['Images'][index]['content']))
                    file.seek(0)
                    file = resize_and_autoorient(file,100,100)
                    p=row[i].add_paragraph()
                    picture = p.add_run()
                    picture.add_picture(file)

                    index +=1
                    

                
            # if item['Images'][0]['content']:
            #     img = item['Images'][0]['content']
            #     file = io.BytesIO(base64.b64decode(img))
            #     file.seek(0)
            #     img = resize_and_autoorient(file, 80,80)
            #     p = doc.add_paragraph()
            #     p.paragraph_format.keep_with_next = True
            #     p.add_run().add_picture(img)
                
                
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Grid Table Light'
        table.style.paragraph_format.keep_with_next=True
        row = table.rows[0].cells
        if "Kommentar" not in item['Items'].keys():item['Items']['Kommentar'] = '-'
        row[0].text = item['Items']['Kommentar']
        row[1].text = item['Items']['Bed_x00f6_mning']['Value']
        p = doc.add_paragraph()
        p.text = item['Items']['Utrustningstyp']['Value']
        p.style = 'small'
        for cell in table.columns[0].cells:
            cell.width = Inches(6)
        for cell in table.columns[1].cells:
            cell.width = Inches(0.4)
    return None

def add_brunnar(doc,js):
    doc.add_paragraph()
    hh = doc.add_heading('Brunnar:', 0)
    hh.style = 'Big heading'
    hh.paragraph_format.keep_with_next=True
    #images = [img['Image'][0] for img in js['Anmärkningar']]
    for i, item in enumerate(js['Brunnar']):
        if item['Items']['{HasAttachments}']:
            table = doc.add_table(rows=0, cols=4)
            index = 0
            table.style.paragraph_format.keep_with_next = True
            while index < len(item['Image']):
                row = table.add_row().cells
                for i in range(0,4):
                    if index == len(item['Image']): continue
                    row[i]
                    file = io.BytesIO(base64.b64decode(item['Image'][index]['content']))
                    file.seek(0)
                    file = resize_and_autoorient(file,100,100)
                    p=row[i].paragraphs[0]
                    row[i].paragraphs[0].paragraph_format.keep_with_next=True
                    picture = p.add_run()
                    picture.add_picture(file)

                    index +=1
                    
                
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Grid Table Light'
        table.style.paragraph_format.keep_with_next=True
        row = table.rows[0].cells
        row[1].text = item['Items']['Anm_x00e4_rkning']
        row[1].paragraphs[0].paragraph_format.keep_with_next=True
        row[0].text = item['Items']['Status']['Value']
        row[0].paragraphs[0].paragraph_format.keep_with_next=True
        p = doc.add_paragraph()
        p.text = "Enligt ordningslagen 3 kap 5§"
        p.style = 'small'
        for cell in table.columns[0].cells:
            cell.width = Inches(6)
        for cell in table.columns[1].cells:
            cell.width = Inches(0.4)
    return None



def change_icon_in_header(doc, js1, icon_file, imagepath):
    """Changes the icon in the header of a word file by replacing a pre-formatted placeholder image."""
    file_content = io.BytesIO(icon_file)
    file_content.seek(0)
    img = Image.open(file_content)
    fb = io.BytesIO()
    img.save(fb, 'PNG')
    fb.seek(0)
    docbyte = io.BytesIO()
    if "mailmerge" in str(type(doc)):
        doc.write(docbyte)
    else: doc.save(docbyte)
    docbyte.seek(0)
    if js1["Certnr"].lower() != 'saknas' and js1['Fitnessbesiktning']:
        doc = add_icon_to_word_file(docbyte.read(), icon_file=fb.read(), imagepath="word/media/image2.png")
    else: doc = add_icon_to_word_file(docbyte.read(), icon_file=fb.read(), imagepath=imagepath)
    doc = docx.Document(doc)
    return doc


def add_page_break(doc):

    # Get the page properties
    section = doc.sections[0]
    page_height = section.page_height
    page_width = section.page_width
    top_margin = section.top_margin
    bottom_margin = section.bottom_margin
    left_margin = section.left_margin
    right_margin = section.right_margin

    # Get the height of the content in the document
    content_height = 0
    for paragraph in doc.paragraphs:
        # Calculate the height of the paragraph using its spacing information
        space_before = paragraph.paragraph_format.space_before
        space_before = space_before.pt if space_before is not None else 0

        space_after = paragraph.paragraph_format.space_after
        space_after = space_after.pt if space_after is not None else 0
        font_size = paragraph.runs[0].font.size.pt if paragraph.runs and paragraph.runs[0].font.size else 0

        line_spacing = paragraph.paragraph_format.line_spacing
        line_spacing = line_spacing if line_spacing is not None else 1.15
        paragraph_height = space_before + space_after + line_spacing * font_size
        content_height += paragraph_height
        section = doc.sections[0]
        page_height = section.page_height

    # Estimate the remaining space on the page
    remaining_space = page_height - top_margin - bottom_margin - content_height


def run_functions(js):
    doc = create_protocol('Funktionskontrolllekplatsdemo',"Lista_lekplats_besiktningsprotokoll",js)
    file = io.BytesIO()
    doc.save(file)
    file.seek(0)

    if "Items" in js.keys():js1 = js['Items']['value'][0]
    else: js1 = js['body']['Items']['value'][0]
    
    #file = compress_word_file(file.getvalue())
    filename="Protokoll_"+str(js1['ID'])+'_'+js1['Title']+'_'+js1['Adress']+'_'+js1['Datum']
    if __name__=='__main__': return doc
    return {"content": base64.b64encode(file.getvalue()).decode('utf-8'), "filename": filename}


if __name__ == '__main__':
    # with open(os.path.join(os.path.dirname(__file__),'sample.json'),'r', encoding="utf-8") as f:
    #     js = json.load(f)
    # doc = create_protocol('Funktionskontrolllekplatsdemo',"Lista_lekplats_besiktningsprotokoll",js)
    test_one=False
    if not test_one:
        jsonpath = os.path.join(os.path.dirname(__file__),'Lekplatsprotokoll_json_filer')
        destpath = os.path.join(os.path.dirname(__file__), 'Testing_word_filer')
        for item in os.listdir(jsonpath):
            if '' in item:    
                filename = os.path.join(jsonpath,item)
                print(item)
                with open(filename,'r', encoding="utf-8") as f:
                    js = json.load(f)
                    doc = run_functions(js)
                doc.save(os.path.join(destpath,item.split('.')[0]+'.docx'))
                
    if test_one:
        with open(os.path.join(os.path.dirname(__file__), 'tt.json'), encoding='utf-8') as f:
            js = json.load(f)
            doc = run_functions(js)
            doc.save(os.path.join(os.path.dirname(__file__),'steg_1.docx'))
        

    ### TODO: