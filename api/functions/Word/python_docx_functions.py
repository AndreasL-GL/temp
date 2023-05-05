import docx


def addCheckbox(para, box_id, name, checked, text_to_add):
  para = para.clear()
  run = para.add_run()
  tag = run._r
  fldchar = docx.oxml.shared.OxmlElement('w:fldChar')
  fldchar.set(docx.oxml.ns.qn('w:fldCharType'), 'begin')

  ffdata = docx.oxml.shared.OxmlElement('w:ffData')
  name = docx.oxml.shared.OxmlElement('w:name')
  name.set(docx.oxml.ns.qn('w:val'), "")
  enabled = docx.oxml.shared.OxmlElement('w:enabled')
  calconexit = docx.oxml.shared.OxmlElement('w:calcOnExit')
  calconexit.set(docx.oxml.ns.qn('w:val'), '0')

  checkbox = docx.oxml.shared.OxmlElement('w:checkBox')
  sizeauto = docx.oxml.shared.OxmlElement('w:sizeAuto')
  default = docx.oxml.shared.OxmlElement('w:default')

  if checked:
    default.set(docx.oxml.ns.qn('w:val'), '1')
  else:
    default.set(docx.oxml.ns.qn('w:val'), '0')

  checkbox.append(sizeauto)
  checkbox.append(default)
  ffdata.append(name)
  ffdata.append(enabled)
  ffdata.append(calconexit)
  ffdata.append(checkbox)
  fldchar.append(ffdata)
  tag.append(fldchar)

  run2 = para.add_run()
  tag2 = run2._r
  start = docx.oxml.shared.OxmlElement('w:bookmarkStart')
  start.set(docx.oxml.ns.qn('w:id'), str(box_id))
  start.set(docx.oxml.ns.qn('w:name'), name)
  tag2.append(start)

  run3 = para.add_run()
  tag3 = run3._r
  instr = docx.oxml.OxmlElement('w:instrText')
  instr.text = 'FORMCHECKBOX'
  tag3.append(instr)

  run4 = para.add_run()
  tag4 = run4._r
  fld2 = docx.oxml.shared.OxmlElement('w:fldChar')
  fld2.set(docx.oxml.ns.qn('w:fldCharType'), 'end')
  tag4.append(fld2)

  run5 = para.add_run()
  tag5 = run5._r
  end = docx.oxml.shared.OxmlElement('w:bookmarkEnd')
  end.set(docx.oxml.ns.qn('w:id'), str(box_id))
  end.set(docx.oxml.ns.qn('w:name'), name)
  tag5.append(end)
  run6 = para.add_run()
  run6.text = text_to_add

  return